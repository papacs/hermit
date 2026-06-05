from __future__ import annotations

import os
import re
import shutil
import stat
from datetime import datetime, timezone
from pathlib import Path, PureWindowsPath
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DEFAULT_WORKSPACE_ROOT = Path(r"C:\HermitWorkspace")
SECURITY_MESSAGE = "SecurityError: 拒绝访问沙箱外文件"
_INVALID_PATH_CHARS = set('\0*?"<>|')


class SecurityError(Exception):
    pass


class InvalidInputError(Exception):
    pass


class SafeDocxProcessor:
    @staticmethod
    def safe_read(file_name: str) -> dict:
        try:
            docx_path = _resolve_docx_path(file_name)
            if not docx_path.exists():
                return _error("FILE_NOT_FOUND", "文件不存在")

            document = Document(docx_path)
            return {
                "ok": True,
                "file": _safe_display_name(file_name),
                "sections": _extract_sections(document),
                "warnings": [],
            }
        except SecurityError:
            return _error("SECURITY_ERROR", SECURITY_MESSAGE)
        except InvalidInputError as exc:
            return _error("INVALID_INPUT", str(exc))
        except Exception:
            return _error("DOCX_OPEN_FAILED", "无法打开文档")

    @staticmethod
    def safe_update_section(
        file_name: str,
        target_heading: str,
        new_text: str,
    ) -> dict:
        try:
            _validate_text_argument(target_heading, "target_heading")
            _validate_text_argument(new_text, "new_text")

            docx_path = _resolve_docx_path(file_name)
            if not docx_path.exists():
                return _error("FILE_NOT_FOUND", "文件不存在")

            try:
                document = Document(docx_path)
            except Exception:
                return _error("DOCX_OPEN_FAILED", "无法打开文档")

            sections = _section_ranges(document)
            matches = [
                section
                for section in sections
                if section["heading"].strip() == target_heading.strip()
            ]
            if not matches:
                return _error("HEADING_NOT_FOUND", "未找到目标标题")
            if len(matches) > 1:
                return _error("HEADING_NOT_UNIQUE", "目标标题不唯一")

            target = matches[0]
            if _has_unsupported_content(document, target["start"], target["end"]):
                return _error("UNSUPPORTED_COMPLEX_CONTENT", "目标区域包含不支持的复杂内容")

            backup_relative = _backup_source(docx_path)
            output_path = _next_revision_path(docx_path)
            changed_count = _replace_section_text(
                document,
                target["start"],
                target["end"],
                new_text,
            )

            try:
                document.save(output_path)
            except Exception:
                return _error("OUTPUT_WRITE_FAILED", "修订版写入失败")

            return {
                "ok": True,
                "file": _safe_display_name(file_name),
                "backup": backup_relative,
                "output": _relative_to_workspace(output_path),
                "targetHeading": target_heading,
                "changedParagraphs": changed_count,
                "warnings": [],
            }
        except SecurityError:
            return _error("SECURITY_ERROR", SECURITY_MESSAGE)
        except InvalidInputError as exc:
            return _error("INVALID_INPUT", str(exc))
        except Exception:
            return _error("OUTPUT_WRITE_FAILED", "修订版写入失败")


def safe_read(file_name: str) -> dict:
    return SafeDocxProcessor.safe_read(file_name)


def safe_update_section(file_name: str, target_heading: str, new_text: str) -> dict:
    return SafeDocxProcessor.safe_update_section(file_name, target_heading, new_text)


def _workspace_root() -> Path:
    configured_root = os.environ.get("HERMIT_WORKSPACE_ROOT")
    root = Path(configured_root) if configured_root else DEFAULT_WORKSPACE_ROOT
    try:
        root.mkdir(parents=True, exist_ok=True)
        (root / ".backup").mkdir(parents=True, exist_ok=True)
        return root.resolve(strict=False)
    except Exception as exc:
        raise InvalidInputError("无法创建安全工作区") from exc


def _resolve_docx_path(file_name: str) -> Path:
    if not isinstance(file_name, str):
        raise InvalidInputError("file_name 必须是字符串")

    raw_name = file_name.strip()
    if not raw_name:
        raise InvalidInputError("file_name 不能为空")
    if any(char in raw_name for char in _INVALID_PATH_CHARS):
        raise SecurityError()
    if ":" in raw_name:
        raise SecurityError()
    if raw_name.startswith("\\\\") or raw_name.startswith("//"):
        raise SecurityError()

    windows_path = PureWindowsPath(raw_name)
    if windows_path.is_absolute():
        raise SecurityError()
    if any(part in ("..", "") for part in windows_path.parts):
        raise SecurityError()
    if windows_path.suffix.lower() != ".docx":
        raise InvalidInputError("仅支持 .docx 文件")

    root = _workspace_root()
    candidate = root.joinpath(*windows_path.parts)
    resolved = candidate.resolve(strict=False)
    if not _is_within_root(resolved, root):
        raise SecurityError()
    if _contains_reparse_point(candidate, root):
        raise SecurityError()
    return resolved


def _is_within_root(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False


def _contains_reparse_point(path: Path, root: Path) -> bool:
    current = root
    relative_parts = []
    try:
        relative_parts = list(path.relative_to(root).parts)
    except ValueError:
        raise SecurityError()

    for part in relative_parts:
        current = current / part
        if not current.exists():
            continue
        if current.is_symlink():
            return True
        attrs = getattr(current.lstat(), "st_file_attributes", 0)
        if attrs & getattr(stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0):
            return True
    return False


def _extract_sections(document: Any) -> list[dict]:
    ranges = _section_ranges(document)
    sections = []
    for item in ranges:
        paragraphs = document.paragraphs[item["start"] + 1 : item["end"]]
        text = "\n".join(paragraph.text for paragraph in paragraphs if paragraph.text)
        sections.append(
            {
                "heading": item["heading"],
                "level": item["level"],
                "paragraphStart": item["start"],
                "paragraphEnd": item["end"] - 1,
                "text": text,
            }
        )
    return sections


def _section_ranges(document: Any) -> list[dict]:
    headings = []
    for index, paragraph in enumerate(document.paragraphs):
        level = _heading_level(paragraph)
        if level and paragraph.text.strip():
            headings.append(
                {
                    "heading": paragraph.text.strip(),
                    "level": level,
                    "start": index,
                }
            )

    sections = []
    for position, heading in enumerate(headings):
        end = len(document.paragraphs)
        for next_heading in headings[position + 1 :]:
            if next_heading["level"] <= heading["level"]:
                end = next_heading["start"]
                break
        sections.append({**heading, "end": end})
    return sections


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = getattr(paragraph.style, "name", "") or ""
    match = re.match(r"^Heading\s+(\d+)$", style_name)
    if match:
        return int(match.group(1))
    return None


def _has_unsupported_content(document: Any, start: int, end: int) -> bool:
    body_elements = list(document._body._element)
    start_element = document.paragraphs[start]._element
    end_element = document.paragraphs[end]._element if end < len(document.paragraphs) else None
    inside_target = False

    for element in body_elements:
        if element is start_element:
            inside_target = True
            continue
        if end_element is not None and element is end_element:
            break
        if not inside_target:
            continue
        if element.tag.lower().endswith("}tbl"):
            return True

    for paragraph in document.paragraphs[start:end]:
        xml = paragraph._element.xml
        if any(token in xml for token in ("<w:drawing", "<w:pict", "<w:commentRangeStart", "<w:ins", "<w:del")):
            return True
    return False


def _replace_section_text(document: Any, start: int, end: int, new_text: str) -> int:
    heading_paragraph = document.paragraphs[start]
    old_paragraphs = document.paragraphs[start + 1 : end]
    body_style = old_paragraphs[0].style if old_paragraphs else None

    for paragraph in old_paragraphs:
        parent = paragraph._element.getparent()
        parent.remove(paragraph._element)

    current = heading_paragraph
    lines = new_text.splitlines() or [""]
    for line in lines:
        current = _insert_paragraph_after(current, line, body_style)
    return len(lines)


def _insert_paragraph_after(paragraph: Paragraph, text: str, style: Any) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def _backup_source(path: Path) -> str:
    root = _workspace_root()
    backup_dir = root / ".backup"
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    backup_name = f"{_safe_stem(path.stem)}_{timestamp}.docx"
    backup_path = backup_dir / backup_name

    try:
        shutil.copy2(path, backup_path)
    except Exception as exc:
        raise InvalidInputError("备份失败") from exc
    return _relative_to_workspace(backup_path)


def _next_revision_path(path: Path) -> Path:
    revision = path.with_name(f"{path.stem}_Hermit修订版{path.suffix}")
    if not revision.exists():
        return revision

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    return path.with_name(f"{path.stem}_Hermit修订版_{timestamp}{path.suffix}")


def _relative_to_workspace(path: Path) -> str:
    root = _workspace_root()
    return path.relative_to(root).as_posix()


def _safe_display_name(file_name: str) -> str:
    return PureWindowsPath(file_name).as_posix()


def _safe_stem(stem: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", stem).strip("._") or "document"


def _validate_text_argument(value: str, name: str) -> None:
    if not isinstance(value, str):
        raise InvalidInputError(f"{name} 必须是字符串")
    if "\0" in value:
        raise InvalidInputError(f"{name} 包含非法字符")


def _error(error_code: str, message: str) -> dict:
    return {
        "ok": False,
        "errorCode": error_code,
        "message": message,
        "warnings": [],
    }

import hashlib
import importlib
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_docx(path: Path, sections: list[tuple[str, list[str]]]) -> None:
    document = Document()
    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    document.save(path)


class SafeDocxProcessorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.workspace = Path(self.temp_dir.name) / "HermitWorkspace"
        os.environ["HERMIT_WORKSPACE_ROOT"] = str(self.workspace)

        import hermit_skills.docx_processor as docx_processor

        self.docx_processor = importlib.reload(docx_processor)
        self.workspace.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        os.environ.pop("HERMIT_WORKSPACE_ROOT", None)
        self.temp_dir.cleanup()

    def test_safe_read_reads_sections_inside_workspace(self) -> None:
        _write_docx(
            self.workspace / "proposal.docx",
            [("Implementation", ["Step one", "Step two"])],
        )

        result = self.docx_processor.safe_read("proposal.docx")

        self.assertTrue(result["ok"])
        self.assertEqual(result["file"], "proposal.docx")
        self.assertEqual(result["sections"][0]["heading"], "Implementation")
        self.assertIn("Step one", result["sections"][0]["text"])

    def test_safe_read_rejects_parent_directory_escape(self) -> None:
        result = self.docx_processor.safe_read("..\\secret.docx")

        self.assertFalse(result["ok"])
        self.assertEqual(result["errorCode"], "SECURITY_ERROR")
        self.assertEqual(result["message"], "SecurityError: 拒绝访问沙箱外文件")
        self.assertNotIn(str(self.workspace), result["message"])

    def test_safe_read_rejects_absolute_path(self) -> None:
        result = self.docx_processor.safe_read(str(self.workspace / "proposal.docx"))

        self.assertFalse(result["ok"])
        self.assertEqual(result["errorCode"], "SECURITY_ERROR")
        self.assertNotIn(str(self.workspace), result["message"])

    def test_safe_read_rejects_non_docx_file(self) -> None:
        (self.workspace / "notes.txt").write_text("not a document", encoding="utf-8")

        result = self.docx_processor.safe_read("notes.txt")

        self.assertFalse(result["ok"])
        self.assertEqual(result["errorCode"], "INVALID_INPUT")

    def test_safe_update_section_creates_backup_and_revision_without_overwriting_source(self) -> None:
        source = self.workspace / "proposal.docx"
        _write_docx(source, [("Implementation", ["Old plan"]), ("Budget", ["100"])])
        before_hash = _sha256(source)

        result = self.docx_processor.safe_update_section(
            "proposal.docx",
            "Implementation",
            "New plan\nSecond paragraph",
        )

        self.assertTrue(result["ok"])
        self.assertEqual(_sha256(source), before_hash)
        self.assertEqual(result["backup"].split("/")[0], ".backup")
        self.assertTrue((self.workspace / result["backup"]).exists())
        output = self.workspace / result["output"]
        self.assertTrue(output.exists())

        updated = Document(output)
        text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
        self.assertIn("New plan", text)
        self.assertIn("Second paragraph", text)
        self.assertIn("Budget", text)
        self.assertNotIn("Old plan", text)

    def test_safe_update_section_fails_when_heading_missing(self) -> None:
        _write_docx(self.workspace / "proposal.docx", [("Implementation", ["Old plan"])])

        result = self.docx_processor.safe_update_section("proposal.docx", "Missing", "New")

        self.assertFalse(result["ok"])
        self.assertEqual(result["errorCode"], "HEADING_NOT_FOUND")

    def test_safe_update_section_fails_when_heading_duplicate(self) -> None:
        _write_docx(
            self.workspace / "proposal.docx",
            [("Implementation", ["A"]), ("Implementation", ["B"])],
        )

        result = self.docx_processor.safe_update_section(
            "proposal.docx",
            "Implementation",
            "New",
        )

        self.assertFalse(result["ok"])
        self.assertEqual(result["errorCode"], "HEADING_NOT_UNIQUE")

    def test_docx_processor_source_does_not_contain_delete_or_exec_apis(self) -> None:
        source = Path(self.docx_processor.__file__).read_text(encoding="utf-8")
        forbidden = [
            "os.remove",
            "os.unlink",
            "Path.unlink",
            "shutil.rmtree",
            "eval(",
            "exec(",
            "compile(",
            "subprocess",
        ]

        for token in forbidden:
            with self.subTest(token=token):
                self.assertNotIn(token, source)


if __name__ == "__main__":
    unittest.main()

